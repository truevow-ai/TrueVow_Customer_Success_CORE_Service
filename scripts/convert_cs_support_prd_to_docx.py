"""
Script to convert CS-Support Service PRD from Markdown to Word (.docx) format.
"""

import os
import re
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("Error: python-docx not installed. Install with: pip install python-docx")
    exit(1)

# Paths
CS_SUPPORT_DIR = r"C:\Users\yasha\OneDrive\Documents\TrueVow\Cursor\TrueVow-CS-Support"
MD_FILE = os.path.join(CS_SUPPORT_DIR, "docs", "CS_SUPPORT_SERVICE_PRD.md")
DOCX_FILE = os.path.join(CS_SUPPORT_DIR, "CS_SUPPORT_SERVICE_PRD.docx")
# Try to use temp file if original is locked
TEMP_DOCX_FILE = os.path.join(CS_SUPPORT_DIR, "CS_SUPPORT_SERVICE_PRD-UPDATED.docx")

def clean_text(text):
    """Remove control characters and NULL bytes that can't be in XML."""
    if not text:
        return text
    # Remove NULL bytes and control characters (except newlines, tabs, carriage returns)
    text = re.sub(r'[\x00-\x08\x0B-\x0C\x0E-\x1F\x7F]', '', text)
    return text

def process_markdown_line(line, doc, in_code_block=False):
    """Process a single markdown line and add to document."""
    line = line.rstrip()
    
    # Skip empty lines
    if not line.strip():
        return False
    
    # Handle code blocks
    if line.startswith('```'):
        return not in_code_block  # Toggle code block state
    
    if in_code_block:
        # Add code as monospace text
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        return True
    
    # Handle headers
    if line.startswith('#'):
        level = len(line) - len(line.lstrip('#'))
        text = line.lstrip('#').strip()
        if text:
            # Remove markdown formatting from header text
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)  # Remove bold
            text = re.sub(r'`(.*?)`', r'\1', text)  # Remove code
            doc.add_heading(clean_text(text), level=min(level, 9))
        return False
    
    # Handle horizontal rules
    if line.strip() == '---':
        doc.add_paragraph('_' * 80)
        return False
    
    # Handle list items
    if line.strip().startswith('- ') or line.strip().startswith('* '):
        text = line.strip()[2:].strip()
        # Process inline formatting
        text = process_inline_formatting(text)
        if text:
            p = doc.add_paragraph(text, style='List Bullet')
        return False
    
    # Handle numbered lists
    if re.match(r'^\d+\.\s+', line.strip()):
        text = re.sub(r'^\d+\.\s+', '', line.strip())
        text = process_inline_formatting(text)
        if text.strip():
            p = doc.add_paragraph(text.strip(), style='List Number')
        return False
    
    # Handle checkboxes
    if re.match(r'^-\s+\[[ xX]\]\s+', line):
        text = re.sub(r'^-\s+\[[ xX]\]\s+', '', line)
        text = process_inline_formatting(text)
        checkbox = '☑' if '[x]' in line.lower() or '[X]' in line else '☐'
        p = doc.add_paragraph(f"{checkbox} {text}", style='List Bullet')
        return False
    
    # Handle bold text (markdown **text**)
    if '**' in line or '*' in line or '`' in line:
        text = process_inline_formatting(line)
        if text.strip():
            p = doc.add_paragraph()
            add_formatted_text(p, text)
        return False
    
    # Regular paragraph
    text = clean_text(line.strip())
    if text:
        # Remove markdown code backticks but keep text
        text = text.replace('`', '')
        p = doc.add_paragraph(text)
    return False

def process_inline_formatting(text):
    """Process inline markdown formatting."""
    # Remove code backticks
    text = re.sub(r'`([^`]+)`', r'\1', text)
    # Remove bold markers (we'll handle formatting in add_formatted_text)
    return clean_text(text)

def add_formatted_text(paragraph, text):
    """Add text with formatting (bold, italic, code)."""
    # Split by ** for bold
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            # Bold text
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*') and len(part) > 2:
            # Italic text (single asterisk, not bold)
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith('`') and part.endswith('`'):
            # Code text
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
        else:
            # Regular text
            if part.strip():
                paragraph.add_run(part)

def convert_markdown_to_docx(md_path, docx_path):
    """Convert markdown file to Word document."""
    print(f"Reading: {os.path.basename(md_path)}")
    
    # Read markdown file
    try:
        with open(md_path, 'r', encoding='utf-8', errors='replace') as f:
            content = f.read()
    except Exception as e:
        print(f"Error reading file: {e}")
        return False
    
    # Create new document
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Process content line by line
    lines = content.split('\n')
    in_code_block = False
    
    for line in lines:
        result = process_markdown_line(line, doc, in_code_block)
        if isinstance(result, bool):
            in_code_block = result
    
    # Add footer with update timestamp
    doc.add_paragraph()
    footer_para = doc.add_paragraph()
    footer_para.add_run(f"Last Updated: {datetime.now().strftime('%B %d, %Y')}").italic = True
    
    # Save document
    try:
        # Try to save directly first
        try:
            doc.save(docx_path)
            print(f"Successfully updated: {os.path.basename(docx_path)}")
            print(f"Location: {docx_path}")
            return True
        except PermissionError:
            # If file is locked, save to alternative file
            import shutil
            alt_path = TEMP_DOCX_FILE
            print(f"Original file is locked (may be open in Word or syncing).")
            print(f"Saving to alternative file: {os.path.basename(alt_path)}")
            doc.save(alt_path)
            print(f"Successfully created: {os.path.basename(alt_path)}")
            print(f"Location: {alt_path}")
            print(f"\nNOTE: Please close the original .docx file, then:")
            print(f"  1. Delete or rename the old CS_SUPPORT_SERVICE_PRD.docx")
            print(f"  2. Rename CS_SUPPORT_SERVICE_PRD-UPDATED.docx to CS_SUPPORT_SERVICE_PRD.docx")
            return True
        
    except Exception as e:
        print(f"Error saving document: {e}")
        import traceback
        traceback.print_exc()
        return False

def main():
    print("=" * 80)
    print("Converting CS-Support Service PRD to Word Document")
    print("=" * 80)
    print()
    
    if not os.path.exists(MD_FILE):
        print(f"Error: Markdown file not found: {MD_FILE}")
        return
    
    if convert_markdown_to_docx(MD_FILE, DOCX_FILE):
        print()
        print("=" * 80)
        print("Conversion completed successfully!")
        print("=" * 80)
    else:
        print()
        print("=" * 80)
        print("Conversion failed!")
        print("=" * 80)

if __name__ == "__main__":
    main()
